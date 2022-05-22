import docx
from os import listdir
from os.path import isfile, join

def read_docx():
    invoices_folder_path: str = '../Invoices/'

    invoices_filenames = [join(invoices_folder_path, f) for f in listdir(invoices_folder_path) if isfile(join(invoices_folder_path, f)) and "Inv" in f]
    print(invoices_filenames)
    print(len(invoices_filenames))
    totals: list[float] = []
    for name in invoices_filenames:
        doc = docx.Document(name)
        # for paragraph in doc.paragraphs:
        #     if "Total" in paragraph.text:
        #         print(paragraph.text)
        
        for table in doc.tables: 
            for row in table.rows: 
                for cell in row.cells:
                    if '£' in cell.text:
                        print("tables next")
                        print(cell.text)
read_docx()
