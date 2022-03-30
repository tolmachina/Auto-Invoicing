from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date
from job_classes import JobInfo
from constants import EVENT_CONCEPT_ADDRESS, ACCOUNT_DATA

class Invoice():
    def __init__(self, job: JobInfo):
        self.job: JobInfo = job
        self.document = Document()
        # heading
        heading_text = "INVOICE " + job.invoice_number
        h = self.document.add_heading(heading_text)
        h.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # first paragraph
        p = self.document.add_paragraph("Bill To:")
        p1 = self.document.add_paragraph()
        p1.add_run(EVENT_CONCEPT_ADDRESS + '\n').italic = True
        # date and num of invoice
        p_date_and_num = self.document.add_paragraph("Invoice Date: " + str(date.today().strftime("%d/%m/%Y")) + '\n')
        p_date_and_num.add_run("Invoice Number: " + str(job.invoice_number)) 
        # text body
        p_service = self.document.add_paragraph("For providing event services for " + job.date_start +  " as agreed with " + job.manager_name + ".\n")
        # table heading
        self.table= self.document.add_table(rows=2, cols=5)  # magic row numbers !!!
        self.table.style = 'Light Shading Accent 1'
        header_cells = self.table.rows[0].cells
        header_cells[0].text = 'Start Date'
        header_cells[1].text = 'End Date'
        header_cells[2].text = 'Location'
        header_cells[3].text = 'Job Number'
        header_cells[4].text = 'Agreed rate' 
        header_cells
        #table body
        first_row = self.table.rows[1].cells
        first_row[0].text = job.date_start
        first_row[1].text = job.date_end
        first_row[2].text = job.location
        first_row[3].text = job.id
        first_row[4].text = job.rate

    def add_table_row(self, new_job):
        new_row = self.table.add_row()
       
        new_row.cells[0].text = new_job.date_start
        new_row.cells[1].text = new_job.date_end
        new_row.cells[2].text = new_job.location
        new_row.cells[3].text = new_job.id
        new_row.cells[4].text = new_job.rate
        
    def finish_doc(self):
        self.total_rate = 0
        for row in self.table.rows[1:]:
            self.total_rate += float(row.cells[4].text[1:])
        new_row = self.table.add_row()
        new_row.cells[3].text = 'Total'
        new_row.cells[4].text = '£' + str(self.total_rate)

        #last paragraph
        self.document.add_paragraph(ACCOUNT_DATA)
        # creating filename and saving file
        self.doc_name = 'Invoice '+ self.job.invoice_number + '.docx'
        self.document.save(self.doc_name)
    
    def get_file_name(self):
        return self.doc_name